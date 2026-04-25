#!/usr/bin/env python3
"""
Add copyright notice to all Waternova .docx files and copy to content/.
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUTHOR = "Federico Blanco Sánchez-Llanos"
YEAR = 2026
COPYRIGHT_TEXT = f"© {YEAR} {AUTHOR}. All rights reserved."

SRC = Path("/root/Waternova")
DST = Path("/root/invinoveritas/content")
DST.mkdir(parents=True, exist_ok=True)

for src_path in sorted(SRC.glob("*.docx")):
    doc = Document(src_path)

    # Check if copyright already inserted
    if doc.paragraphs and COPYRIGHT_TEXT in doc.paragraphs[0].text:
        print(f"  skip (already has copyright): {src_path.name}")
        shutil.copy2(src_path, DST / src_path.name)
        continue

    # Insert copyright as first paragraph
    first_para = doc.paragraphs[0]._element
    new_para = doc.add_paragraph()
    new_para.text = COPYRIGHT_TEXT
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.runs[0]
    run.font.size = Pt(9)
    run.font.italic = True

    # Move it to the top
    doc.paragraphs[0]._element.addprevious(new_para._element)

    dst_path = DST / src_path.name
    doc.save(dst_path)
    print(f"  ✓ {src_path.name}")

print(f"\nDone. {len(list(DST.glob('*.docx')))} files in {DST}")
